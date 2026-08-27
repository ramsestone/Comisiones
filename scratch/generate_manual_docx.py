# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style:
            b_element = OxmlElement(f'w:{border_name}')
            b_element.set(qn('w:val'), border_style.get('val', 'single'))
            b_element.set(qn('w:sz'), str(border_style.get('sz', 4)))
            b_element.set(qn('w:space'), '0')
            b_element.set(qn('w:color'), border_style.get('color', 'CCCCCC'))
            tcBorders.append(b_element)
        else:
            b_element = OxmlElement(f'w:{border_name}')
            b_element.set(qn('w:val'), 'none')
            tcBorders.append(b_element)
    tcPr.append(tcBorders)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    h.paragraph_format.space_before = Pt(14 if level == 1 else (10 if level == 2 else 6))
    h.paragraph_format.space_after = Pt(4)
    run = h.runs[0]
    if level == 1:
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 60, 60) # Dark Teal
    elif level == 2:
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(22, 101, 52) # Dark Green
    elif level == 3:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate
    return h

def add_bullet_point(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_bold = p.add_run(bold_prefix)
    r_bold.bold = True
    r_bold.font.color.rgb = RGBColor(15, 23, 42)
    r_text = p.add_run(text)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_numbered_step(doc, number_str, bold_title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    
    r_num = p.add_run(f"{number_str} ")
    r_num.bold = True
    r_num.font.color.rgb = RGBColor(22, 101, 52)
    
    if bold_title:
        r_title = p.add_run(f"{bold_title} ")
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(15, 23, 42)
        
    r_text = p.add_run(text)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_callout(doc, title, text, style_type='info'):
    # Colors based on style
    if style_type == 'warning':
        bg_color = "FEF3C7" # light amber
        border_color = "D97706"
        title_color = RGBColor(180, 83, 9)
    elif style_type == 'danger':
        bg_color = "FEE2E2" # light red
        border_color = "DC2626"
        title_color = RGBColor(185, 28, 28)
    else: # info/note
        bg_color = "ECFDF5" # light emerald
        border_color = "059669"
        title_color = RGBColor(4, 120, 87)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    set_cell_borders(cell, 
                     left={'val': 'single', 'sz': 24, 'color': border_color},
                     top={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                     right={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                     bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"📌 {title}\n" if title else "")
    r_title.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = title_color
    
    r_text = p.add_run(text)
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    
    # Empty paragraph after table for spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def create_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A3A") # Dark Teal
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        set_cell_borders(hdr_cells[i], 
                         top={'val': 'single', 'sz': 6, 'color': '0F2828'},
                         bottom={'val': 'single', 'sz': 12, 'color': '166534'},
                         left={'val': 'single', 'sz': 4, 'color': '2E5A5A'},
                         right={'val': 'single', 'sz': 4, 'color': '2E5A5A'})
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (i > 1 and len(title) < 15) else WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    # Data rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_hex = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_hex)
            set_cell_margins(row_cells[c_idx], top=90, bottom=90, left=140, right=140)
            set_cell_borders(row_cells[c_idx], 
                             top={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                             bottom={'val': 'single', 'sz': 4, 'color': 'CBD5E1' if r_idx == len(data)-1 else 'E2E8F0'},
                             left={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                             right={'val': 'single', 'sz': 4, 'color': 'E2E8F0'})
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            # Align checkmarks or numbers to center
            if str(val).strip() in ['✓', '—', 'Sí', 'No', 'Completo', 'Pendiente']:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(9)
                if str(val).strip() == '✓':
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(22, 163, 74)
                elif str(val).strip() == '—':
                    r.font.color.rgb = RGBColor(148, 163, 184)
                else:
                    r.font.color.rgb = RGBColor(51, 65, 85)
                    
    # Column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
                
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)
    return table

def build_manual_document():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles config
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(51, 65, 85)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # ── COVER / TITLE HEADER ──────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(4)
    r_sys = p_title.add_run("SISTEMA NOVU\n")
    r_sys.font.name = 'Arial Black'
    r_sys.font.size = Pt(24)
    r_sys.font.color.rgb = RGBColor(15, 60, 60)
    
    r_sub = p_title.add_run("MANUAL DE USUARIO FUNCIONAL\n")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(15)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(22, 101, 52)
    
    r_desc = p_title.add_run("Plataforma de Gestión Integral de Comisiones Inmobiliarias, Cartera y KPIs\nVersión 1.1 — Documento Operativo y de Capacitación Empresarial")
    r_desc.font.size = Pt(9.5)
    r_desc.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── 1. INTRODUCCIÓN ──────────────────────────────────────────────
    add_styled_heading(doc, "1. Introducción", level=1)
    
    add_styled_heading(doc, "1.1 Objetivo del Sistema", level=2)
    doc.add_paragraph(
        "El Sistema NOVU es una solución empresarial de alta gama diseñada para centralizar, auditar y agilizar "
        "todo el ciclo operativo de comisiones por comercialización inmobiliaria. La plataforma abarca desde el "
        "registro inicial de la venta (Contratos, Escrituras, Bonos y Cancelaciones), la verificación individual "
        "de porcentajes por parte de los asesores y gerentes, la validación y autorización de pago por Dirección "
        "y Administración, hasta el control financiero de préstamos y abonos (Cartera de Deudores) y el análisis "
        "ejecutivo mediante Dashboards de Indicadores Clave de Desempeño (KPIs)."
    )

    add_styled_heading(doc, "1.2 Alcance Funcional", level=2)
    add_bullet_point(doc, "Control de Acceso y Seguridad: ", "Autenticación segura por roles con permisos segmentados.")
    add_bullet_point(doc, "Sincronización con ERP Maestro (V10 / Enkontrol): ", "Integración de catálogos maestros de compañías, desarrollos, ubicaciones y expedientes de clientes.")
    add_bullet_point(doc, "Gestión Integral de Comisiones: ", "Captura, cálculo automático, esquema Multipunto (MP), verificación colaborativa, pase a aprobación y marcado de pago.")
    add_bullet_point(doc, "Historial Consolidado por Propiedad: ", "Trazabilidad por fases de vivienda (Contrato, Escritura y Bono) y registro de escrituración en 1 clic.")
    add_bullet_point(doc, "Cartera de Deudores y Wallet: ", "Administración independiente de deudas activas (préstamos/penalizaciones) y comisiones/saldos a favor con aplicación de abonos.")
    add_bullet_point(doc, "Dashboard Ejecutivo de KPIs: ", "Analítica de volumen vendido, comisiones totales, absorción de inventario y rankings comerciales.")
    add_bullet_point(doc, "Administración de Usuarios: ", "Control de colaboradores, perfiles, asignación de gerentes a asesores y gestión de datos de contacto.")

    add_styled_heading(doc, "1.3 Perfiles de Usuario", level=2)
    add_bullet_point(doc, "Administrador: ", "Acceso total al sistema, altas/bajas de usuarios, autorización de pagos finales, sincronización V10, auditoría de cartera y reportes globales.")
    add_bullet_point(doc, "Director / Directora: ", "Revisión ejecutiva, aprobación y solicitud de correcciones a comisiones verificadas, consulta de KPIs globales, ranking de gerentes y supervisión de cartera.")
    add_bullet_point(doc, "Gerente: ", "Captura de comisiones, verificación de montos de su equipo comercial, consulta de historial y seguimiento a cartera de deudores.")
    add_bullet_point(doc, "Asesor: ", "Consulta de comisiones personales asignadas, verificación individual de montos y revisión de su historial y estado de cuenta.")

    add_styled_heading(doc, "1.4 Requisitos para Utilizar el Sistema", level=2)
    add_bullet_point(doc, "Navegador Web: ", "Google Chrome, Microsoft Edge, Mozilla Firefox o Safari (versiones modernas actualizadas).")
    add_bullet_point(doc, "Conectividad: ", "Acceso a la red corporativa para garantizar sincronización directa con los servicios de base de datos V10/Enkontrol.")
    add_bullet_point(doc, "Cuenta de Usuario: ", "Credenciales individuales asignadas formalmente por el Administrador.")

    # ── 2. ACCESO AL SISTEMA ─────────────────────────────────────────
    add_styled_heading(doc, "2. Acceso al Sistema", level=1)
    doc.add_paragraph(
        "Para iniciar sus actividades en la plataforma, acceda a la dirección web oficial del Sistema NOVU. "
        "En la pantalla de bienvenida se presentará el formulario de autenticación segura."
    )
    
    add_callout(doc, "Captura pendiente de agregar", "Figura 1. Pantalla de Inicio de Sesión (Login) con visualizador de contraseña interactivo.")

    add_styled_heading(doc, "2.1 Procedimiento de Inicio de Sesión", level=2)
    add_numbered_step(doc, "1.", "Nombre de Usuario:", "Capture su nombre de usuario asignado en el primer campo.")
    add_numbered_step(doc, "2.", "Contraseña:", "Ingrese su contraseña de acceso.")
    add_numbered_step(doc, "3.", "Verificar Clave:", "(Opcional) Haga clic sobre el icono del ojo para visualizar temporalmente los caracteres y confirmar su correcta escritura.")
    add_numbered_step(doc, "4.", "Iniciar Sesión:", "Haga clic en el botón 'Iniciar Sesión'. El sistema validará sus credenciales y lo redirigirá al panel principal.")

    add_styled_heading(doc, "2.2 Mensajes y Validaciones de Acceso", level=2)
    create_table(doc, 
        ["Situación / Error", "Mensaje del Sistema", "Causa y Solución"],
        [
            ["Campos Vacíos", "Usuario y contraseña son requeridos", "Debe capturar ambos campos obligatorios."],
            ["Datos Incorrectos", "Usuario o contraseña incorrectos", "Verifique que no existan errores tipográficos o mayúsculas activadas."],
            ["Cuenta Inactiva", "Usuario o contraseña incorrectos", "Si su cuenta fue dada de baja administrativamente, contacte al Administrador."]
        ],
        [1.8, 2.3, 2.4]
    )

    add_styled_heading(doc, "2.3 Cierre de Sesión Seguro", level=2)
    doc.add_paragraph(
        "Para proteger la confidencialidad de los datos, al concluir su jornada laboral haga clic en el botón 'Salir' "
        "ubicado en la esquina superior derecha del encabezado. Se cerrará la sesión de forma segura y se destruirá la "
        "cookie de autenticación."
    )

    # ── 3. INTERFAZ PRINCIPAL ─────────────────────────────────────────
    add_styled_heading(doc, "3. Interfaz Principal y Navegación", level=1)
    doc.add_paragraph(
        "La interfaz del sistema está estructurada en un entorno ergonómico de alta productividad con dos áreas clave:"
    )
    add_bullet_point(doc, "Menú Lateral (Sidebar): ", "Permite saltar entre los módulos autorizados: Comisiones, Nueva Comisión, Historial, KPIs, Cartera de Deudores, Gestión de Usuarios y Registrar Usuario.")
    add_bullet_point(doc, "Barra Superior (Header): ", "Contiene el título de la pantalla activa, el botón de Sincronización Manual V10 (Admin), la campana de Notificaciones con badge dinámico, la tarjeta de perfil y el botón de salida.")
    add_bullet_point(doc, "Buzón de Notificaciones: ", "Alerta sobre comisiones que requieren su verificación, aprobaciones emitidas o comisiones enviadas a corrección. Al hacer clic sobre cualquier notificación se redirige de forma automática a la comisión respectiva.")

    # ── 4. BANDEJA DE COMISIONES ─────────────────────────────────────
    add_styled_heading(doc, "4. Módulo de Comisiones (Bandeja Operativa)", level=1)
    doc.add_paragraph(
        "Es el núcleo operativo donde se monitorean todas las comisiones en curso. Dependiendo de su perfil, el título "
        "se adaptará a 'Mis comisiones' (Asesores/Gerentes) o 'Todas las comisiones' (Dirección/Administración)."
    )

    add_callout(doc, "Captura pendiente de agregar", "Figura 2. Bandeja de Comisiones con filtros avanzados, etiquetas de estado y tarjetas de participantes.")

    add_styled_heading(doc, "4.1 Filtros y Herramientas de Auditoría", level=2)
    add_bullet_point(doc, "Buscador Textual: ", "Busca coincidencias inmediatas por nombre de cliente, lote/ubicación, desarrollo, empresa o colaborador participante.")
    add_bullet_point(doc, "Filtro por Estado: ", "Permite aislar comisiones en: Pendiente Verificación, Verificada, Pendiente Aprobación, Aprobada, Pendiente Pago, Pagada, Corrección y Cancelada.")
    add_bullet_point(doc, "Selectores de Empresa, Desarrollo y Concepto: ", "Menús desplegables generados dinámicamente según los registros existentes.")
    add_bullet_point(doc, "Rango de Fechas (Desde / Hasta): ", "Permite acotar el listado a un periodo de fechas de registro específico.")
    add_bullet_point(doc, "Exportación a CSV: ", "(Visible para Director y Administrador) Descarga el reporte auditado con 12 columnas incluyendo desarrollo, cliente, vendedor, precio, comisiones desglosadas y tipo de penalización.")

    add_styled_heading(doc, "4.2 Ciclo de Vida y Estados de una Comisión", level=2)
    doc.add_paragraph(
        "1. Pendiente Verificación (Amarillo): La comisión recién creada espera que asesor y gerente confirmen sus montos.\n"
        "2. Pendiente Aprobación (Azul / Amarillo): Todos los participantes han verificado; la comisión está lista para revisión de Dirección.\n"
        "3. Corrección (Rojo): La Dirección o Administración devolvió la comisión indicando inconsistencias para su ajuste.\n"
        "4. Pendiente Pago (Azul): Aprobada por Dirección; en espera de dispersión por Administración.\n"
        "5. Pagada (Verde): Liquidada formalmente en su totalidad.\n"
        "6. Cancelada (Gris): Operación rescidida con desglose de penalización."
    )

    add_styled_heading(doc, "4.3 Procedimientos Clave en la Bandeja", level=2)
    add_numbered_step(doc, "A.", "Verificación Individual de Participantes:", 
                      "El asesor o gerente localiza su tarjeta, identifica su nombre con la etiqueta '(tú)' y hace clic en 'Verificar'. El botón cambiará a 'Verificado'. Al completar todas las firmas de la tarjeta, el sistema notifica a Dirección.")
    add_numbered_step(doc, "B.", "Aprobación por Dirección:", 
                      "El Director revisa la tarjeta en estado 'Pendiente Aprobación' y hace clic en 'Aprobar'. La comisión pasa a 'Pendiente Pago' y se notifica al Administrador.")
    add_numbered_step(doc, "C.", "Solicitud de Corrección:", 
                      "Si la Dirección o Administración detecta un error, pulsa 'Corregir', escribe los motivos en el modal emergente y confirma. La comisión pasa a estado 'Corrección' con un banner explicativo.")
    add_numbered_step(doc, "D.", "Marcado de Pago:", 
                      "El Administrador presiona 'Marcar como pagada' tras realizar la transferencia bancaria o aplicación correspondiente.")

    # ── 5. REGISTRO Y EDICIÓN DE COMISIONES ──────────────────────────
    add_styled_heading(doc, "5. Registro y Edición de Comisiones", level=1)
    doc.add_paragraph(
        "Formulario estructurado para la creación y modificación de comisiones comerciales con integración en tiempo real "
        "a los catálogos maestros de Enkontrol/V10."
    )

    add_callout(doc, "Captura pendiente de agregar", "Figura 3. Formulario de Captura de Comisión y Modal de Confirmación Previa al Guardado.")

    add_styled_heading(doc, "5.1 Estructura de Campos del Formulario", level=2)
    create_table(doc,
        ["Campo", "Descripción Funcional", "Obligatorio", "Regla de Negocio"],
        [
            ["Compañía", "Razón social del desarrollo inmobiliario", "Sí", "Selector con búsqueda."],
            ["Desarrollo", "Proyecto o condominio habitacional", "Sí", "Filtrado por la compañía."],
            ["Ubicación", "Unidad / Vivienda / Extra (Bodega)", "Sí", "Carga el precio de venta en automático."],
            ["Expediente", "Identificador de cliente y contrato", "Sí", "Formato [ID] - [Nombre Cliente]."],
            ["Concepto", "Tipo de hito comercial", "Sí", "Contrato, Escritura, Cancelación o Bono."],
            ["Destino Penalización", "¿A quién aplica la deducción?", "Condicional", "Exclusivo si Concepto = Cancelación (Cliente vs Miembros)."],
            ["Fecha Operación", "Fecha de firma de contrato/escritura", "Sí", "Selector de calendario."],
            ["Precio de Venta", "Monto pactado de la vivienda", "Sí", "Cargado automáticamente desde V10."],
            ["Gerentes (Filas 1-3)", "Gerentes comisionistas", "Sí", "Selector de personal y porcentaje."],
            ["Asesores (Filas 1-3)", "Asesores comisionistas", "Sí", "Selector de personal y porcentaje."],
            ["Multipunto (MP)", "Esquema de venta compartida", "No", "Casilla de verificación (*checkbox*)."]
        ],
        [1.5, 2.4, 0.9, 1.7]
    )

    add_styled_heading(doc, "5.2 Modal de Previsualización y Confirmación", level=2)
    doc.add_paragraph(
        "Al hacer clic en 'Guardar Comisión', el sistema no envía los datos a ciegas; despliega un modal oscuro "
        "con el desglose exacto en pesos ($) y porcentajes (%) de cada participante, la compañía y la ubicación. "
        "El usuario puede pulsar 'Editar' para regresar al formulario o 'Confirmar y Guardar' para emitir la comisión."
    )

    # ── 6. HISTORIAL DE COMISIONES ────────────────────────────────────
    add_styled_heading(doc, "6. Historial Consolidado por Propiedad", level=1)
    doc.add_paragraph(
        "Organiza la información agrupándola por Unidad Inmobiliaria (Ubicación), permitiendo visualizar en un solo bloque "
        "todas las comisiones que ha devengado una misma vivienda a lo largo del tiempo (fases de Contrato, Escritura y Bono)."
    )

    add_callout(doc, "Captura pendiente de agregar", "Figura 4. Vista de Historial por Ubicación con Acordeón de Fases Desplegado.")

    add_styled_heading(doc, "6.1 Funcionalidad 'Registrar Escritura desde Contrato'", level=2)
    doc.add_paragraph(
        "Cuando una vivienda ya tiene registrada su comisión de Contrato y llega el momento de la escrituración:\n"
        "1. Ingrese a Historial y busque la propiedad.\n"
        "2. Abra el bloque de la vivienda y en la columna 'Escritura' haga clic en 'Registrar comisión de escritura'.\n"
        "3. El formulario se abrirá con la compañía, desarrollo, lote, cliente, precio y participantes precargados y protegidos.\n"
        "4. Capture únicamente la fecha de escritura y confirme. Esto previene errores de captura y discrepancias entre fases."
    )

    # ── 7. DASHBOARD DE KPIS ──────────────────────────────────────────
    add_styled_heading(doc, "7. Dashboard de KPIs y Analítica Comercial", level=1)
    doc.add_paragraph(
        "Proporciona a la Dirección y Gerencia comercial un centro de mando visual para la toma de decisiones estratégicas "
        "con métricas actualizadas en tiempo real."
    )

    add_callout(doc, "Captura pendiente de agregar", "Figura 5. Dashboard de Métricas KPIs, Rankings Comerciales y Tabla de Desarrollos.")

    add_styled_heading(doc, "7.1 Tarjetas de Indicadores Ejecutivos", level=2)
    add_bullet_point(doc, "Volumen Total Vendido: ", "Monto global acumulado en pesos de operaciones cerradas exitosamente.")
    add_bullet_point(doc, "Comisiones Generadas: ", "Importe total de comisiones distribuidas al equipo comercial.")
    add_bullet_point(doc, "Cierres Exitosos: ", "Total de contratos y escrituras formalizados en el periodo.")
    add_bullet_point(doc, "Cancelaciones y Tasa (%): ", "Conteo de cancelaciones y porcentaje relativo sobre el total de ventas.")
    add_bullet_point(doc, "Comisión Promedio: ", "Ticket promedio de comisión por operación cerrada.")

    add_styled_heading(doc, "7.2 Tabla de Desarrollos y Absorción de Inventario", level=2)
    doc.add_paragraph(
        "Tabla con 9 columnas que compara las viviendas totales del desarrollo, las vendidas en V10 (ERP maestro), "
        "las vendidas en Sistema NOVU, el porcentaje de absorción de mercado y los montos vendidos/cancelados. "
        "Cuenta con ordenamiento dinámico de 3 clics (Descendente -> Ascendente -> Original) y botón de Exportar a CSV."
    )

    # ── 8. CARTERA DE DEUDORES ────────────────────────────────────────
    add_styled_heading(doc, "8. Cartera de Deudores y Control de Saldos", level=1)
    doc.add_paragraph(
        "Módulo financiero para la gestión de cuentas corrientes de asesores y gerentes. Mantiene una separación estricta "
        "entre las deudas activas (préstamos o penalizaciones por cancelación) y los saldos positivos ganados."
    )

    add_callout(doc, "Captura pendiente de agregar", "Figura 6. Ficha de Cartera de Deudores con herramientas de cobro/abono.")

    add_styled_heading(doc, "8.1 Operaciones Financieras Disponibles", level=2)
    add_bullet_point(doc, "Registro de Préstamo / Cargo: ", "Permite agregar una deuda activa al colaborador indicando monto y concepto.")
    add_bullet_point(doc, "Aplicar Abono a Deuda: ", "Permite liquidar deuda descontando de las comisiones activas o saldo a favor del colaborador. El sistema valida que el abono no exceda la deuda ni el saldo disponible.")
    add_bullet_point(doc, "Historial de Movimientos: ", "Auditoría completa de abonos, cargos, comisiones y créditos por usuario.")

    # ── 9. GESTIÓN DE USUARIOS ────────────────────────────────────────
    add_styled_heading(doc, "9. Gestión de Usuarios y Accesos", level=1)
    doc.add_paragraph(
        "Módulo reservado para el Administrador. Permite dar de alta colaboradores, definir roles, asignar gerentes "
        "a asesores, editar datos de emergencia, actualizar fotos de perfil y dar de baja lógica cuentas."
    )

    add_callout(doc, "Captura pendiente de agregar", "Figura 7. Módulo de Gestión de Usuarios y Formulario de Alta.")

    # ── 10. MATRIZ DE PERMISOS ────────────────────────────────────────
    add_styled_heading(doc, "10. Matriz de Permisos por Rol", level=1)
    create_table(doc,
        ["Funcionalidad / Módulo", "Administrador", "Director", "Gerente", "Asesor"],
        [
            ["Acceso a Login y Mi Perfil", "✓", "✓", "✓", "✓"],
            ["Ver 'Mis Comisiones'", "✓", "✓", "✓", "✓"],
            ["Ver 'Todas las Comisiones' (Global)", "✓", "✓", "—", "—"],
            ["Verificar Participación Individual", "✓", "✓", "✓", "✓"],
            ["Registrar Nueva Comisión", "✓", "✓", "✓", "—"],
            ["Editar / Eliminar Comisión propia", "✓", "✓", "✓", "—"],
            ["Aprobar Comisiones (Pase a Pago)", "✓ (Dev)", "✓", "—", "—"],
            ["Solicitar Corrección con Comentarios", "✓", "✓", "—", "—"],
            ["Marcar Comisión como Pagada", "✓", "—", "—", "—"],
            ["Ver Comisión de Dirección", "✓", "✓", "—", "—"],
            ["Registrar Escritura desde Contrato", "✓", "✓", "✓", "—"],
            ["Historial Consolidado de Fases", "✓", "✓", "✓", "✓"],
            ["Dashboard de KPIs y Analítica", "✓", "✓", "✓", "—"],
            ["Ranking 'Top Gerentes'", "✓", "✓", "—", "—"],
            ["Exportar Tablas a CSV", "✓", "✓", "—", "—"],
            ["Cartera de Deudores (Wallet)", "✓", "✓", "✓", "—"],
            ["Sincronizar Base de Datos V10", "✓", "✓", "—", "—"],
            ["Crear y Modificar Usuarios", "✓", "—", "—", "—"]
        ],
        [2.3, 1.1, 1.1, 1.0, 1.0]
    )

    # ── 11. CASOS DE USO FRECUENTES ──────────────────────────────────
    add_styled_heading(doc, "11. Casos de Uso Frecuentes ('¿Cómo hago...?')", level=1)
    
    add_bullet_point(doc, "¿Cómo registro una comisión nueva?: ", 
                     "Vaya a 'Nueva Comisión', seleccione Compañía -> Desarrollo -> Ubicación -> Expediente, confirme fecha y precio, asigne Gerente/Asesor con sus porcentajes y pulse 'Guardar Comisión'.")
    add_bullet_point(doc, "¿Cómo registro la escritura de una vivienda ya comisionada en contrato?: ", 
                     "Vaya a 'Historial', abra la ubicación deseada y en la fase de Escritura pulse 'Registrar comisión de escritura'. Capture la fecha y confirme.")
    add_bullet_point(doc, "¿Cómo verifico mi comisión?: ", 
                     "En 'Comisiones', busque la tarjeta amarilla 'Pendiente Verificación', localice su nombre con la etiqueta '(tú)' y pulse 'Verificar'.")
    add_bullet_point(doc, "¿Cómo atiendo una comisión en 'Corrección'?: ", 
                     "Abra la tarjeta roja, lea el banner de observaciones de Dirección, pulse 'Editar', realice los cambios solicitados y guarde nuevamente.")
    add_bullet_point(doc, "¿Cómo abono a la deuda de un asesor?: ", 
                     "En 'Cartera de Deudores', abra la ficha del asesor, capture el monto a pagar en el campo de abono y presione 'Aplicar Abono a Deuda'.")

    # ── 12. MENSAJES Y SOLUCIÓN DE PROBLEMAS ─────────────────────────
    add_styled_heading(doc, "12. Mensajes del Sistema y Solución de Problemas", level=1)
    create_table(doc,
        ["Problema / Mensaje", "Causa Probable", "Acción Correctiva"],
        [
            ["'Por favor selecciona una compañía/desarrollo'", "Campos obligatorios incompletos.", "Complete los selectores marcados con asterisco rojo."],
            ["'El monto a abonar excede el saldo disponible'", "Se intentó abonar más dinero del que el asesor tiene a favor.", "Ajuste el importe a una cantidad igual o menor al saldo positivo."],
            ["'Error al conectar con servidor V10'", "Pérdida temporal de enlace con el ERP Enkontrol.", "Verifique la conexión de red corporativa o reintente la sincronización."],
            ["Selectores de desarrollo vacíos", "No se ha seleccionado la compañía matriz.", "Seleccione primero la compañía correspondiente para filtrar desarrollos."],
            ["No aparece 'Nueva Comisión' o 'Usuarios'", "Su rol actual no posee permisos para ese módulo.", "Consulte la matriz de permisos o contacte a su Administrador."]
        ],
        [2.1, 2.1, 2.3]
    )

    # ── 13. GLOSARIO DE TÉRMINOS ─────────────────────────────────────
    add_styled_heading(doc, "13. Glosario de Términos", level=1)
    add_bullet_point(doc, "Absorción (%): ", "Indicador comercial que mide el avance porcentual de ventas sobre el total proyectado de viviendas en un desarrollo.")
    add_bullet_point(doc, "Comisión Multipunto (MP): ", "Modalidad de venta compartida entre distintos puntos de venta o colaboradores foráneos.")
    add_bullet_point(doc, "Deuda Activa: ", "Saldo deudor pendiente de pago atribuible a un colaborador por préstamos corporativos o cancelaciones.")
    add_bullet_point(doc, "Enkontrol / V10: ", "Sistema ERP empresarial de gestión inmobiliaria conectado con NOVU para la sincronización de catálogos maestros.")
    add_bullet_point(doc, "Penalización al Cliente: ", "Cancelación donde el importe penalizado al cliente cubre el costo de la operación sin descontar comisiones a la fuerza de ventas.")
    add_bullet_point(doc, "Penalización a Miembros: ", "Cancelación donde se genera un cargo deudor a asesores y gerentes para recuperar comisiones pagadas de una venta no concretada.")
    add_bullet_point(doc, "Saldo a Favor: ", "Importe positivo disponible del colaborador derivado de comisiones ganadas o créditos directos.")

    # ── 14. MATRIZ DE COBERTURA ──────────────────────────────────────
    add_styled_heading(doc, "14. Matriz Final de Cobertura de Funcionalidades", level=1)
    create_table(doc,
        ["Módulo", "Funcionalidad Documentada", "Doc.", "Captura", "Permisos", "Flujo", "Estado"],
        [
            ["Autenticación", "Login, Toggle Password, Logout, Sesión", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Comisiones", "Bandeja, Filtros, Verificación de firmas", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Comisiones", "Aprobación Dirección, Corrección, Pago", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Comisiones", "Edición, Eliminación y Exportación CSV", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Registro", "Cascadas, Autoselección V10, Multipunto", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Registro", "Cancelaciones, Modal de Confirmación", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Historial", "Agrupación por ubicación, Fases de venta", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Historial", "Registro de Escritura directa en 1 clic", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["KPIs", "Métricas, Rankings Asesores/Gerentes", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["KPIs", "Tabla Desarrollos 3 clics, Gráficas Chart.js", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Cartera", "Resumen saldos, Préstamos y Abonos", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["Usuarios", "Altas, roles, asignación de gerentes, bajas", "✓", "Pendiente", "✓", "✓", "Completo"],
            ["General", "Notificaciones en tiempo real, Sinc V10", "✓", "Pendiente", "✓", "✓", "Completo"]
        ],
        [1.1, 2.3, 0.5, 0.8, 0.6, 0.5, 0.7]
    )

    # ── 15. PENDIENTES DE DOCUMENTACIÓN ──────────────────────────────
    add_styled_heading(doc, "15. Pendientes de Documentación", level=1)
    add_bullet_point(doc, "Capturas de Pantalla Definitivas: ", "Se dejaron señalizadas las figuras para insertar las capturas finales del entorno en producción.")
    add_bullet_point(doc, "Tabulador de Bonos Especiales: ", "⚠️ Pendiente de confirmar con Dirección Comercial las condiciones de elegibilidad para la fase de 'Bono'.")
    add_bullet_point(doc, "Plazos Reglamentarios de Verificación: ", "⚠️ Pendiente de confirmar si existe un tiempo límite en días hábiles para la verificación de comisiones antes del corte.")

    # Guardar documento
    output_path = r"c:\Users\Ramses Obregon\Desktop\Proyectos\Comission-Managment\Manual_de_Usuario_NOVU.docx"
    doc.save(output_path)
    print(f"Documento generado exitosamente en: {output_path}")

if __name__ == "__main__":
    build_manual_document()
